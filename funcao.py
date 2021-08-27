from pydub import AudioSegment
import speech_recognition as sr
from os import remove, chdir
from docx import Document

class Acao:
    def __init__(self,video_caminho):
        self.reco = sr.Recognizer() #Declarando um reconhecedor do Recognize Speech 
        self.video_caminho = video_caminho
        self.tempo = AudioSegment.from_file(str(self.video_caminho), "mp4").duration_seconds
        self.num_pacotes = int(self.tempo/150)
        self.resto = (self.tempo - int(self.tempo))
    
    def reconhecimento_propriamente(self,audio):
        with sr.AudioFile(audio) as som_bruto:
            #self.reco.adjust_for_ambient_noise(som_bruto) #Torna o som mais limpo
            som_limpo = self.reco.record(som_bruto) #Armazena o som limpo na memória 
            #Agora cria-se uma tupla responsável por armazenar todos os dados relativos a análise do áudio
            resposta = 'TEXTO NULO'
            try:
                resposta = str(self.reco.recognize_google(som_limpo,language='pt-BR')) #Aqui efetivamente tenta-se usar a API da Google pra reconhecer as palavras no áudio
                
            except sr.RequestError: #Caso haja um erro de comunicação com a API
                resposta = "API não disponível. Por favor, verifique sua conexão com a internet!"
                
            except sr.UnknownValueError: #Caso a API retorne um erro de reconhecimento/processamento do áudio 
                resposta = "O áudio está incompreensível"
            
            return resposta

    def fracionador(self,pasta_destino): #Fracionador do vídeo em várias faixas de áudio
        i=1
        doc = Document()
        doc.add_heading('TRANSCRIÇÃO DO TEXTO',0)
        print(pasta_destino)
        for i in range(1,self.num_pacotes):
            faixa = AudioSegment.from_file(str(self.video_caminho), "mp4") #Redefine o AudioSegmente a cada vez que o loop inicia
            faixa[(i-1)*150000:(i*150000)].export("audio.wav",format="wav") #Segmenta o áudio em pacotes de 2,5 minutos
            doc.add_paragraph(self.reconhecimento_propriamente('audio.wav')) #A cada ciclo, adiciona um parágrafo ao arquivo .docx
            remove('audio.wav')
            
        faixa = AudioSegment.from_file(str(self.video_caminho), "mp4")
        faixa[-self.resto*1000:].export("audio.wav",format="wav") #O que sobrou do áudio (um pacote menor que 2,5 minutos) é transcrito
        doc.add_paragraph(self.reconhecimento_propriamente('audio.wav')) #Adiciona a leitura desse último parágrafo
        doc.add_page_break()
        remove('audio.wav')
        chdir(pasta_destino)
        doc.save('Transcrição.docx')
        